import io
import re
from typing import Optional


SUPPORTED_EXTENSIONS = {'pdf', 'docx', 'doc', 'bpmn'}


def _extract_bpmn(file_bytes: bytes) -> str:
    return file_bytes.decode('utf-8', errors='replace').strip()


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())
    return "\n\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def _detect_embedded_bpmn(text: str) -> Optional[str]:
    """Return BPMN XML if found embedded in the extracted text."""
    pattern = (
        r'(<\?xml[^>]*\?>[\s\S]*?<\/bpmn:definitions>)'
        r'|(<bpmn:definitions[\s\S]*?<\/bpmn:definitions>)'
        r'|(<definitions[\s\S]*?<\/definitions>)'
    )
    match = re.search(pattern, text, re.DOTALL)
    if match:
        return match.group(0).strip()
    return None


def extract_document(file_bytes: bytes, filename: str) -> dict:
    """
    Extract content from a BPMN, PDF, or Word document.

    Returns:
        {
            "text": str,            # full extracted text (empty for .bpmn)
            "has_bpmn": bool,       # always True for .bpmn files
            "bpmn_xml": str|None,   # BPMN XML (the file itself for .bpmn, or detected inside doc)
            "filename": str,
            "pages": int|None,      # page count for PDFs
            "char_count": int,
            "file_type": str        # "bpmn" | "pdf" | "word"
        }
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    if ext not in SUPPORTED_EXTENSIONS:
        return {"error": f"Unsupported file type '.{ext}'. Please upload a BPMN (.bpmn), PDF (.pdf), or Word (.docx, .doc) file."}

    # ── .bpmn file: treat the file itself as the BPMN XML to update ──────────
    if ext == 'bpmn':
        try:
            bpmn_xml = _extract_bpmn(file_bytes)
        except Exception as e:
            return {"error": f"Failed to read BPMN file: {str(e)}"}

        if not bpmn_xml.strip():
            return {"error": "The BPMN file appears to be empty."}

        # Basic validity check
        if not any(tag in bpmn_xml for tag in ('<bpmn:definitions', '<definitions', '<?xml')):
            return {"error": "The file does not appear to be a valid BPMN XML file."}

        return {
            "text": "",
            "has_bpmn": True,
            "bpmn_xml": bpmn_xml,
            "filename": filename,
            "pages": None,
            "char_count": len(bpmn_xml),
            "file_type": "bpmn",
        }

    # ── PDF / Word: extract text, check for embedded BPMN ───────────────────
    try:
        if ext == 'pdf':
            import pdfplumber  # trigger ImportError early if missing
            text = _extract_pdf(file_bytes)
            page_count = None
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages)
            file_type = "pdf"
        else:
            text = _extract_docx(file_bytes)
            page_count = None
            file_type = "word"
    except ImportError:
        missing = 'pdfplumber' if ext == 'pdf' else 'python-docx'
        return {"error": f"Missing dependency: {missing}. Run: pip install {missing}"}
    except Exception as e:
        return {"error": f"Failed to extract document: {str(e)}"}

    if not text.strip():
        return {"error": "The document appears to be empty or could not be read."}

    bpmn_xml = _detect_embedded_bpmn(text)

    return {
        "text": text,
        "has_bpmn": bpmn_xml is not None,
        "bpmn_xml": bpmn_xml,
        "filename": filename,
        "pages": page_count,
        "char_count": len(text),
        "file_type": file_type,
    }
